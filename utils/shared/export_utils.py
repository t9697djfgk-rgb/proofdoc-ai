"""Export / download utilities for legal tool results."""
from __future__ import annotations
import json
import re
import streamlit as st


def download_txt(label: str, text: str, filename: str, key: str | None = None) -> None:
    st.download_button(
        label=label,
        data=text,
        file_name=filename,
        mime="text/plain",
        use_container_width=True,
        key=key,
    )


def download_json(label: str, data: dict | list, filename: str, key: str | None = None) -> None:
    st.download_button(
        label=label,
        data=json.dumps(data, indent=2, ensure_ascii=False),
        file_name=filename,
        mime="application/json",
        use_container_width=True,
        key=key,
    )


def download_md(label: str, text: str, filename: str, key: str | None = None) -> None:
    st.download_button(
        label=label,
        data=text,
        file_name=filename,
        mime="text/markdown",
        use_container_width=True,
        key=key,
    )


_BODY_FONT = "Book Antiqua"


def _style_run(run, bold: bool = False, size=None, color=None) -> None:
    """Apply standard font styling to a run."""
    from docx.shared import Pt
    _body_size = Pt(12)
    run.font.name = _BODY_FONT
    run.font.size = size or _body_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _style_heading(p, level: int) -> None:
    """Apply Book Antiqua bold styling to all runs in a heading paragraph."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    sizes = {1: Pt(14), 2: Pt(13), 3: Pt(12)}
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = _BODY_FONT
        run.font.size = sizes.get(level, Pt(12))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)


def _build_docx(text: str, title: str = "") -> bytes:
    """Convert text (plain or markdown) to a properly styled DOCX. Returns bytes."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Firm header
    hdr_para = doc.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr_para.add_run("eLawFirm")
    _style_run(run, bold=True, size=Pt(16), color=RGBColor(0x1a, 0x27, 0x44))

    if title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        _style_run(tr, bold=True, size=Pt(13), color=RGBColor(0x1a, 0x27, 0x44))

    # Horizontal rule after header
    hr_p = doc.add_paragraph()
    pPr = hr_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A84C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()  # spacer

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Markdown headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _style_heading(p, 3)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            _style_heading(p, 2)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _style_heading(p, 1)
            i += 1
            continue

        # Underline-style heading: next line is === or ---
        if i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            if next_stripped and all(c in "=" for c in next_stripped) and len(next_stripped) >= 3:
                p = doc.add_heading(stripped, level=1)
                _style_heading(p, 1)
                i += 2
                continue
            if next_stripped and all(c in "-" for c in next_stripped) and len(next_stripped) >= 3:
                p = doc.add_heading(stripped, level=2)
                _style_heading(p, 2)
                i += 2
                continue

        # Horizontal rule lines
        if stripped in ("---", "===", "***") or (len(set(stripped)) == 1 and stripped[0] in "-=*" and len(stripped) >= 4):
            doc.add_paragraph()
            i += 1
            continue

        # ALL-CAPS short section headers
        is_section_header = (
            stripped == stripped.upper()
            and len(stripped) >= 4
            and len(stripped) <= 60
            and not stripped.startswith("-")
            and any(c.isalpha() for c in stripped)
        )
        if is_section_header:
            p = doc.add_heading(stripped.rstrip(":"), level=2)
            _style_heading(p, 2)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline_formatted_run(p, content)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline_formatted_run(p, content)
            i += 1
            continue

        # Standalone bold line
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r2 = p.add_run(stripped[2:-2])
            _style_run(r2, bold=True)
            i += 1
            continue

        # Normal paragraph with possible inline formatting
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        _add_inline_formatted_run(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_formatted_run(para, text: str) -> None:
    """Add text to a paragraph handling **bold** and _italic_ inline."""
    parts = re.split(r"(\*\*.*?\*\*|_.*?_)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            _style_run(run, bold=True)
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = para.add_run(part[1:-1])
            _style_run(run)
            run.font.italic = True
        else:
            run = para.add_run(part)
            _style_run(run)


def dict_to_markdown(data: dict | list, title: str = "") -> str:
    """Convert any AI result dict/list to readable markdown text for DOCX export."""
    lines: list[str] = []
    if title:
        lines += [f"# {title}", ""]

    def _fmt_val(key: str, val) -> None:
        readable = key.replace("_", " ").title()
        if isinstance(val, list):
            if not val:
                return
            lines.append(f"## {readable}")
            for item in val:
                if isinstance(item, dict):
                    parts = []
                    for k, v in item.items():
                        if v and not isinstance(v, (dict, list)):
                            parts.append(f"**{k.replace('_',' ').title()}:** {v}")
                    if parts:
                        lines.append("- " + " · ".join(parts))
                elif item:
                    lines.append(f"- {item}")
            lines.append("")
        elif isinstance(val, dict):
            lines.append(f"## {readable}")
            for k, v in val.items():
                if v and not isinstance(v, (dict, list)):
                    lines.append(f"- **{k.replace('_',' ').title()}:** {v}")
            lines.append("")
        elif val and not isinstance(val, bool):
            lines.append(f"**{readable}:** {val}")
            lines.append("")

    if isinstance(data, dict):
        # Priority fields first
        priority = ["executive_summary", "summary", "overview", "introduction",
                    "held", "facts", "ratio_decidendi", "conclusion", "client_advice"]
        for k in priority:
            if k in data:
                _fmt_val(k, data[k])
        for k, v in data.items():
            if k not in priority:
                _fmt_val(k, v)
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                parts = [f"**{k.replace('_',' ').title()}:** {v}"
                         for k, v in item.items() if v and not isinstance(v, (dict, list))]
                lines.append("- " + " · ".join(parts))
            else:
                lines.append(f"- {item}")

    return "\n".join(lines)


def download_docx(label: str, text: str, filename: str,
                  title: str = "", key: str | None = None) -> None:
    """Create a properly styled Word document and offer it as a download."""
    try:
        data = _build_docx(text, title=title)
    except Exception:
        import io
        from docx import Document
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
    st.download_button(
        label=label,
        data=data,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key=key,
    )


def download_docx_from_dict(label: str, result: dict | list, filename: str,
                             title: str = "", key: str | None = None) -> None:
    """Convert a result dict to markdown, then to a styled DOCX download."""
    text = dict_to_markdown(result, title=title)
    download_docx(label, text, filename, title=title, key=key)


def download_pdf(label: str, text: str, filename: str, title: str = "", key: str | None = None) -> None:
    """Render plain text as a PDF using fpdf2 and offer it as a download."""
    import io, textwrap
    from fpdf import FPDF

    def _s(s: str) -> str:
        # fpdf2 built-in fonts only cover Latin-1; replace anything outside that range
        return s.encode("latin-1", errors="replace").decode("latin-1")

    _title = _s(title)

    class _PDF(FPDF):
        def header(self):
            if _title:
                self.set_font("Helvetica", "B", 13)
                self.cell(self.epw, 10, _title, align="C", new_x="LMARGIN", new_y="NEXT")
                self.ln(2)

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    w = pdf.epw  # effective width: paper - left margin - right margin (170 mm for A4)

    for raw in text.split("\n"):
        line = _s(raw)
        stripped = line.strip()
        # Pre-wrap very long lines so a single token can never exceed the cell width
        chunks = textwrap.wrap(line, width=120) if len(line) > 120 else [line]
        for chunk in (chunks or [""]):
            try:
                if stripped.startswith("---") or stripped.startswith("==="):
                    pdf.set_draw_color(200, 168, 76)
                    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + w, pdf.get_y())
                    pdf.ln(3)
                    break  # only draw the rule once per original line
                elif (stripped and stripped == stripped.upper()
                      and len(stripped) > 3
                      and all(c.isalpha() or c.isspace() for c in stripped)):
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.multi_cell(w, 6, chunk)
                    pdf.set_font("Helvetica", size=10)
                else:
                    pdf.multi_cell(w, 5.5, chunk if chunk else " ")
            except Exception:
                pdf.ln(5.5)  # skip unrenderable chunk, preserve spacing

    buf = io.BytesIO(pdf.output())
    st.download_button(
        label=label,
        data=buf.read(),
        file_name=filename,
        mime="application/pdf",
        use_container_width=True,
        key=key,
    )


def action_row(
    text_to_download: str,
    base_filename: str,
    report_data: dict | list | None = None,
    reset_keys: list[str] | None = None,
    key_prefix: str = "",
    doc_title: str = "",
) -> None:
    """
    Renders a standard 4-button action row:
    Download .txt | Export report .json | Download .docx | Reset
    """
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        download_txt("📥 Download (.txt)", text_to_download,
                     f"{base_filename}.txt", key=f"{key_prefix}_dl_txt")
    with c2:
        if report_data is not None:
            download_json("📊 Export Report (.json)", report_data,
                          f"{base_filename}_report.json", key=f"{key_prefix}_dl_json")
        else:
            st.write("")
    with c3:
        download_docx("📝 Download (.docx)", text_to_download,
                      f"{base_filename}.docx", title=doc_title, key=f"{key_prefix}_dl_docx")
    with c4:
        if st.button("🔄 Reset", use_container_width=True, key=f"{key_prefix}_reset"):
            if reset_keys:
                for k in reset_keys:
                    st.session_state.pop(k, None)
            st.rerun()


def save_to_matter_ui(text: str, doc_title: str, key_prefix: str) -> None:
    """Expander that saves an AI-generated document to the Document Library under a matter."""
    with st.expander("💾 Save to Matter / Document Library"):
        import utils.database as _db
        matters = []
        try:
            matters = _db.list_matters(status="Active")
        except Exception:
            pass
        if not matters:
            st.info("No active matters found. Create a matter in the Matters page first.")
            return
        matter_opts = {
            f"{m.get('ref', '')} — {(m.get('title') or '')[:40]}": m["id"]
            for m in matters
        }
        sel = st.selectbox("Select Matter", list(matter_opts.keys()),
                           key=f"{key_prefix}_stm_sel")
        doc_name = st.text_input("Document Name", value=doc_title,
                                  key=f"{key_prefix}_stm_name")
        if st.button("💾 Save Document", key=f"{key_prefix}_stm_btn", type="primary",
                     use_container_width=True):
            mid = matter_opts[sel]
            saved = _db.add_document(
                name=doc_name.strip() or doc_title,
                matter_id=mid,
                file_type="text/plain",
                file_size=len((text or "").encode()),
                description=(text or "")[:10000],
                visibility="internal",
            )
            if saved:
                matter_label = sel.split("—")[0].strip()
                st.success(f"✅ Saved to **{matter_label}**. View it in Documents.")
                st.session_state.pop(f"{key_prefix}_stm_saved", None)
            else:
                st.error("Save failed — check Supabase connection.")
