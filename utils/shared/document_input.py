"""Shared document input component and text extraction utilities."""
from __future__ import annotations
import io
import streamlit as st


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")

    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts)

    if ext == "pdf":
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages)

    raise ValueError(f"Unsupported file type: .{ext}. Use PDF, DOCX, or TXT.")


def document_input_ui(
    key_prefix: str,
    accept: list[str] | None = None,
    paste_height: int = 200,
    paste_placeholder: str = "Paste your document text here…",
    label: str = "Upload or paste your document",
) -> str:
    """
    Renders a two-column upload + paste UI.
    Returns the extracted text (empty string if nothing provided).
    """
    if accept is None:
        accept = ["pdf", "docx", "txt"]

    left, right = st.columns([1, 1], gap="large")

    with left:
        st.markdown(f"**📎 Upload** ({', '.join(f'.{e}' for e in accept).upper()})")
        uploaded = st.file_uploader(
            label, type=accept, key=f"{key_prefix}_upload", label_visibility="collapsed"
        )
        if uploaded:
            st.info(f"📄 **{uploaded.name}** · {uploaded.size / 1024:.1f} KB")

    with right:
        st.markdown("**✏️ Or Paste Text**")
        pasted = st.text_area(
            "Paste text",
            height=paste_height,
            placeholder=paste_placeholder,
            key=f"{key_prefix}_paste",
            label_visibility="collapsed",
        )

    text = ""
    if uploaded:
        try:
            text = extract_text(uploaded.read(), uploaded.name)
        except Exception as exc:
            st.error(f"Could not read file: {exc}")
    elif pasted and pasted.strip():
        text = pasted.strip()

    if text:
        char_count = len(text)
        word_count = len(text.split())
        st.caption(f"📊 {word_count:,} words · {char_count:,} characters")

    return text


def two_document_input_ui(key_prefix: str, accept: list[str] | None = None) -> tuple[str, str]:
    """
    Two-document input (original + revised) for comparison tools.
    Returns (original_text, revised_text).
    """
    if accept is None:
        accept = ["pdf", "docx", "txt"]

    st.markdown("**Original Version**")
    original = document_input_ui(f"{key_prefix}_orig", accept=accept,
                                  paste_placeholder="Paste original document text here…")

    st.divider()

    st.markdown("**Revised Version**")
    revised = document_input_ui(f"{key_prefix}_rev", accept=accept,
                                 paste_placeholder="Paste revised document text here…")

    return original, revised
