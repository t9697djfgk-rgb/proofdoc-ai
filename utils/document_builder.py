from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


class DocumentBuilder:
    def build_word_document(
        self, extracted_data: dict, original_filename: str, output_path: str
    ) -> str:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        self._add_header(doc, extracted_data, original_filename)
        doc.add_page_break()

        for block in extracted_data.get("content_blocks", []):
            self._add_block(doc, block)

        for i, tbl in enumerate(extracted_data.get("tables", [])):
            doc.add_heading(f"Table {i + 1}", level=3)
            self._add_table(doc, tbl)
            doc.add_paragraph()

        if extracted_data.get("signatures"):
            self._add_signatures(doc, extracted_data["signatures"])

        doc.save(output_path)
        return output_path

    def _add_header(self, doc, data, filename):
        h = doc.add_heading("PROOFDOC AI — RECONSTRUCTED DOCUMENT", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl = doc.add_table(rows=6, cols=2)
        tbl.style = "Table Grid"
        rows = [
            ("Original File:", filename),
            ("Document Type:", data.get("document_type", "Unknown").replace("_", " ").title()),
            ("Pages Processed:", str(data.get("total_pages", 1))),
            ("AI Confidence:", f"{data.get('avg_confidence', 0):.1f}%"),
            ("Reconstruction Date:", datetime.now().strftime("%B %d, %Y at %H:%M")),
            ("Powered By:", "ProofDoc AI (Claude Opus 4.7)"),
        ]
        for i, (label, value) in enumerate(rows):
            tbl.cell(i, 0).text = label
            tbl.cell(i, 1).text = value
            runs = tbl.cell(i, 0).paragraphs[0].runs
            if runs:
                runs[0].bold = True

        legal = {k: v for k, v in data.get("legal_elements", {}).items() if v}
        if legal:
            doc.add_paragraph()
            doc.add_heading("Detected Legal Elements:", level=2)
            for key, values in legal.items():
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{key.upper()}: ").bold = True
                preview = ", ".join(str(v) for v in values[:10])
                if len(values) > 10:
                    preview += f" (+{len(values) - 10} more)"
                p.add_run(preview)

        if data.get("warnings"):
            doc.add_paragraph()
            doc.add_heading("Review Required:", level=2)
            for w in data["warnings"]:
                doc.add_paragraph(w, style="List Bullet")

    def _add_block(self, doc, block):
        text = block.get("text", "").strip()
        if not text:
            return

        btype = block.get("type", "paragraph")
        align = _ALIGN.get(block.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)

        if btype == "heading":
            level = min(int(block.get("level", 1)), 9)
            p = doc.add_heading(text, level=level)
            p.alignment = align
        elif btype == "numbered_list":
            doc.add_paragraph(text, style="List Number")
        elif btype == "signature_block":
            p = doc.add_paragraph()
            p.add_run(text).italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif btype in ("header", "footer"):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            p = doc.add_paragraph(text)
            p.alignment = align
            fmt = block.get("formatting", "normal")
            if fmt and p.runs:
                run = p.runs[0]
                run.bold = "bold" in fmt
                run.italic = "italic" in fmt
                run.underline = "underline" in fmt

    def _add_table(self, doc, table_data):
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if not headers and not rows:
            return

        total_rows = len(rows) + (1 if headers else 0)
        cols = max(len(headers), max((len(r) for r in rows), default=0))
        if total_rows == 0 or cols == 0:
            return

        tbl = doc.add_table(rows=total_rows, cols=cols)
        tbl.style = "Table Grid"

        row_idx = 0
        if headers:
            for j, h in enumerate(headers[:cols]):
                cell = tbl.cell(0, j)
                cell.text = str(h)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            row_idx = 1

        for i, row in enumerate(rows):
            for j, val in enumerate(row[:cols]):
                tbl.cell(row_idx + i, j).text = str(val)

    def _add_signatures(self, doc, signatures):
        doc.add_page_break()
        doc.add_heading("SIGNATURE BLOCKS", level=1)
        for i, sig in enumerate(signatures, 1):
            doc.add_heading(f"Signature Block {i}", level=2)
            for s in sig.get("signatories", []):
                p = doc.add_paragraph()
                p.add_run("Party: ").bold = True
                p.add_run(str(s))
            if sig.get("has_stamp"):
                doc.add_paragraph("[Official stamp / seal detected]")
            if sig.get("has_date_line"):
                doc.add_paragraph("Date: ___________________________")
            doc.add_paragraph("_" * 40)
            doc.add_paragraph("Signature")
            doc.add_paragraph()
