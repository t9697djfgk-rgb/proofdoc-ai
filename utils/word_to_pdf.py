from fpdf import FPDF
from docx import Document as DocxDocument


def _safe(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def convert_word_to_pdf(docx_path: str, pdf_path: str) -> None:
    doc = DocxDocument(docx_path)

    pdf = FPDF(format="A4")
    pdf.set_margins(22, 22, 22)
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.add_page()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(3)
            continue

        style = para.style.name
        if "Heading 1" in style:
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(30, 58, 95)
            pdf.ln(4)
            pdf.multi_cell(0, 9, _safe(text))
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
        elif "Heading 2" in style:
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(30, 58, 95)
            pdf.ln(3)
            pdf.multi_cell(0, 8, _safe(text))
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)
        elif "Heading" in style:
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 7, _safe(text))
        else:
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            is_italic = any(run.italic for run in para.runs if run.text.strip())
            style_str = ""
            if is_bold and is_italic:
                style_str = "BI"
            elif is_bold:
                style_str = "B"
            elif is_italic:
                style_str = "I"
            pdf.set_font("Helvetica", style_str, 10)
            pdf.multi_cell(0, 6, _safe(text))
            pdf.ln(1)

    for table in doc.tables:
        pdf.ln(4)
        col_count = max(len(row.cells) for row in table.rows) if table.rows else 1
        col_w = 166 / col_count
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if i == 0:
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_fill_color(30, 58, 95)
                pdf.set_text_color(255, 255, 255)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.set_fill_color(248, 249, 250) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(0, 0, 0)
            for cell in cells:
                pdf.cell(col_w, 7, _safe(cell[:40]), border=1, fill=True)
            pdf.ln()
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

    pdf.output(pdf_path)
