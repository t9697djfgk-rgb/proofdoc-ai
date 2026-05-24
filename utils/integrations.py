"""
integrations.py — All integration helpers for eLawFirm.
Uses only `requests` for HTTP calls (no additional SDKs required).
All functions catch exceptions and return {"error": str(e)} or empty list on failure.
"""
from __future__ import annotations

import base64
import hashlib
import io
import json
import os
import re
import time
import uuid
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta, timezone
from typing import Any
from urllib.parse import urlencode, quote

import requests

# ---------------------------------------------------------------------------
# GOOGLE DRIVE
# ---------------------------------------------------------------------------

_GDRIVE_AUTH_URL = "https://accounts.google.com/o/oauth2/v2/auth"
_GDRIVE_TOKEN_URL = "https://oauth2.googleapis.com/token"
_GDRIVE_API = "https://www.googleapis.com/drive/v3"
_GDRIVE_UPLOAD = "https://www.googleapis.com/upload/drive/v3"
_GDRIVE_SCOPE = "https://www.googleapis.com/auth/drive"


def gdrive_auth_url(client_id: str, client_secret: str) -> str:  # noqa: ARG001
    """Return the OAuth2 authorization URL for Google Drive (OOB flow)."""
    try:
        params = {
            "client_id": client_id,
            "redirect_uri": "urn:ietf:wg:oauth:2.0:oob",
            "response_type": "code",
            "scope": _GDRIVE_SCOPE,
            "access_type": "offline",
            "prompt": "consent",
        }
        return f"{_GDRIVE_AUTH_URL}?{urlencode(params)}"
    except Exception as e:
        return f"error:{e}"


def gdrive_exchange_code(client_id: str, client_secret: str, code: str) -> dict:
    """Exchange authorization code for access/refresh tokens."""
    try:
        r = requests.post(
            _GDRIVE_TOKEN_URL,
            data={
                "client_id": client_id,
                "client_secret": client_secret,
                "code": code.strip(),
                "redirect_uri": "urn:ietf:wg:oauth:2.0:oob",
                "grant_type": "authorization_code",
            },
            timeout=15,
        )
        data = r.json()
        if "access_token" not in data:
            return {"error": data.get("error_description", data.get("error", "Token exchange failed"))}
        return {
            "access_token": data.get("access_token", ""),
            "refresh_token": data.get("refresh_token", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def gdrive_upload(
    access_token: str,
    filename: str,
    content_bytes: bytes,
    mimetype: str,
    folder_id: str = "",
) -> dict:
    """Upload a file to Google Drive. Returns {id, url, error?}."""
    try:
        metadata: dict[str, Any] = {"name": filename}
        if folder_id:
            metadata["parents"] = [folder_id]

        meta_part = json.dumps(metadata).encode()
        boundary = "boundary_gdrive_upload"
        body = (
            f"--{boundary}\r\n"
            f"Content-Type: application/json; charset=UTF-8\r\n\r\n"
        ).encode() + meta_part + (
            f"\r\n--{boundary}\r\n"
            f"Content-Type: {mimetype}\r\n\r\n"
        ).encode() + content_bytes + f"\r\n--{boundary}--".encode()

        r = requests.post(
            f"{_GDRIVE_UPLOAD}/files?uploadType=multipart",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": f"multipart/related; boundary={boundary}",
            },
            data=body,
            timeout=60,
        )
        data = r.json()
        if "id" not in data:
            return {"error": data.get("error", {}).get("message", "Upload failed")}
        file_id = data["id"]
        return {
            "id": file_id,
            "url": f"https://drive.google.com/file/d/{file_id}/view",
        }
    except Exception as e:
        return {"error": str(e)}


def gdrive_list(
    access_token: str,
    query: str = "",
    folder_id: str = "",
) -> list[dict]:
    """List files in Google Drive. Returns list of {id, name, mimeType, modifiedTime, size}."""
    try:
        q_parts = ["trashed=false"]
        if folder_id:
            q_parts.append(f"'{folder_id}' in parents")
        if query:
            q_parts.append(f"name contains '{query}'")
        q_str = " and ".join(q_parts)

        r = requests.get(
            f"{_GDRIVE_API}/files",
            headers={"Authorization": f"Bearer {access_token}"},
            params={
                "q": q_str,
                "fields": "files(id,name,mimeType,modifiedTime,size)",
                "pageSize": 50,
                "orderBy": "modifiedTime desc",
            },
            timeout=15,
        )
        data = r.json()
        if "files" not in data:
            return [{"error": data.get("error", {}).get("message", "List failed")}]
        return data["files"]
    except Exception as e:
        return [{"error": str(e)}]


def gdrive_download(access_token: str, file_id: str) -> bytes | None:
    """Download a file from Google Drive. Returns bytes or None on failure."""
    try:
        r = requests.get(
            f"{_GDRIVE_API}/files/{file_id}?alt=media",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=60,
        )
        if r.status_code == 200:
            return r.content
        return None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# ONEDRIVE (Microsoft Device Code Flow)
# ---------------------------------------------------------------------------

_MS_DEVICE_CODE_URL = "https://login.microsoftonline.com/{tenant}/oauth2/v2.0/devicecode"
_MS_TOKEN_URL = "https://login.microsoftonline.com/{tenant}/oauth2/v2.0/token"
_GRAPH_API = "https://graph.microsoft.com/v1.0"
_ONEDRIVE_SCOPE = "Files.ReadWrite offline_access"


def onedrive_device_code(client_id: str, tenant: str = "common") -> dict:
    """Start Device Code Flow for OneDrive. Returns {device_code, user_code, verification_uri, interval, error?}."""
    try:
        r = requests.post(
            _MS_DEVICE_CODE_URL.format(tenant=tenant),
            data={
                "client_id": client_id,
                "scope": _ONEDRIVE_SCOPE,
            },
            timeout=15,
        )
        data = r.json()
        if "device_code" not in data:
            return {"error": data.get("error_description", data.get("error", "Device code request failed"))}
        return {
            "device_code": data["device_code"],
            "user_code": data.get("user_code", ""),
            "verification_uri": data.get("verification_uri", "https://microsoft.com/devicelogin"),
            "interval": data.get("interval", 5),
        }
    except Exception as e:
        return {"error": str(e)}


def onedrive_poll_token(
    client_id: str,
    device_code: str,
    tenant: str = "common",
) -> dict:
    """Poll for token after user completes device code auth. Returns {access_token, refresh_token, error?}."""
    try:
        r = requests.post(
            _MS_TOKEN_URL.format(tenant=tenant),
            data={
                "client_id": client_id,
                "grant_type": "urn:ietf:params:oauth:grant-type:device_code",
                "device_code": device_code,
            },
            timeout=15,
        )
        data = r.json()
        if "access_token" not in data:
            err = data.get("error", "")
            if err == "authorization_pending":
                return {"error": "authorization_pending"}
            return {"error": data.get("error_description", data.get("error", "Token poll failed"))}
        return {
            "access_token": data["access_token"],
            "refresh_token": data.get("refresh_token", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def onedrive_upload(
    access_token: str,
    filename: str,
    content_bytes: bytes,
    folder_path: str = "/",
) -> dict:
    """Upload a file to OneDrive. Returns {id, url, error?}."""
    try:
        clean_folder = folder_path.strip("/")
        if clean_folder:
            url = f"{_GRAPH_API}/me/drive/root:/{clean_folder}/{filename}:/content"
        else:
            url = f"{_GRAPH_API}/me/drive/root:/{filename}:/content"

        r = requests.put(
            url,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/octet-stream",
            },
            data=content_bytes,
            timeout=60,
        )
        if r.status_code not in (200, 201):
            try:
                err = r.json().get("error", {}).get("message", "Upload failed")
            except Exception:
                err = f"HTTP {r.status_code}"
            return {"error": err}
        data = r.json()
        return {
            "id": data.get("id", ""),
            "url": data.get("webUrl", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def onedrive_list(access_token: str, folder_path: str = "/") -> list[dict]:
    """List files in OneDrive folder. Returns list of {id, name, size, lastModifiedDateTime}."""
    try:
        clean = folder_path.strip("/")
        if clean:
            url = f"{_GRAPH_API}/me/drive/root:/{clean}:/children"
        else:
            url = f"{_GRAPH_API}/me/drive/root/children"

        r = requests.get(
            url,
            headers={"Authorization": f"Bearer {access_token}"},
            params={"$select": "id,name,size,lastModifiedDateTime", "$top": 50},
            timeout=15,
        )
        data = r.json()
        if "value" not in data:
            return [{"error": data.get("error", {}).get("message", "List failed")}]
        return [
            {
                "id": f.get("id", ""),
                "name": f.get("name", ""),
                "size": f.get("size", 0),
                "lastModifiedDateTime": f.get("lastModifiedDateTime", ""),
            }
            for f in data["value"]
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# DROPBOX
# ---------------------------------------------------------------------------

_DROPBOX_UPLOAD_URL = "https://content.dropboxapi.com/2/files/upload"
_DROPBOX_LIST_URL = "https://api.dropboxapi.com/2/files/list_folder"
_DROPBOX_SHARE_URL = "https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings"
_DROPBOX_EXISTING_LINK = "https://api.dropboxapi.com/2/sharing/list_shared_links"


def dropbox_upload(
    access_token: str,
    filename: str,
    content_bytes: bytes,
    folder_path: str = "/",
) -> dict:
    """Upload a file to Dropbox. Returns {path, url, error?}."""
    try:
        folder = ("/" + folder_path.strip("/")).rstrip("/")
        full_path = f"{folder}/{filename}" if folder != "/" else f"/{filename}"

        api_arg = json.dumps({
            "path": full_path,
            "mode": "add",
            "autorename": True,
            "mute": False,
        })
        r = requests.post(
            _DROPBOX_UPLOAD_URL,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Dropbox-API-Arg": api_arg,
                "Content-Type": "application/octet-stream",
            },
            data=content_bytes,
            timeout=60,
        )
        data = r.json()
        if "path_display" not in data:
            return {"error": data.get("error_summary", data.get("error", "Upload failed"))}
        path = data["path_display"]
        share = dropbox_get_link(access_token, path)
        return {"path": path, "url": share or ""}
    except Exception as e:
        return {"error": str(e)}


def dropbox_list(access_token: str, folder_path: str = "") -> list[dict]:
    """List files in a Dropbox folder. Returns list of {id, name, size, .tag}."""
    try:
        r = requests.post(
            _DROPBOX_LIST_URL,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json={"path": folder_path, "recursive": False},
            timeout=15,
        )
        data = r.json()
        if "entries" not in data:
            return [{"error": data.get("error_summary", data.get("error", "List failed"))}]
        return [
            {
                "id": e.get("id", ""),
                "name": e.get("name", ""),
                "size": e.get("size", 0),
                ".tag": e.get(".tag", "file"),
            }
            for e in data["entries"]
        ]
    except Exception as e:
        return [{"error": str(e)}]


def dropbox_get_link(access_token: str, file_path: str) -> str | None:
    """Get or create a shared link for a Dropbox file."""
    try:
        r = requests.post(
            _DROPBOX_SHARE_URL,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json={"path": file_path, "settings": {}},
            timeout=15,
        )
        data = r.json()
        if "url" in data:
            return data["url"]
        # May already exist — get existing
        r2 = requests.post(
            _DROPBOX_EXISTING_LINK,
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json={"path": file_path},
            timeout=15,
        )
        d2 = r2.json()
        links = d2.get("links", [])
        if links:
            return links[0].get("url")
        return None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# DOCUSIGN
# ---------------------------------------------------------------------------

def docusign_send_envelope(
    access_token: str,
    account_id: str,
    base_url: str,
    doc_name: str,
    doc_base64: str,
    signer_email: str,
    signer_name: str,
    subject: str = "Please sign",
) -> dict:
    """Send a DocuSign envelope for signature. Returns {envelope_id, status, error?}."""
    try:
        base_url = base_url.rstrip("/")
        envelope = {
            "emailSubject": subject,
            "documents": [{
                "documentBase64": doc_base64,
                "name": doc_name,
                "fileExtension": "pdf",
                "documentId": "1",
            }],
            "recipients": {
                "signers": [{
                    "email": signer_email,
                    "name": signer_name,
                    "recipientId": "1",
                    "routingOrder": "1",
                    "tabs": {
                        "signHereTabs": [{
                            "anchorString": "/sn1/",
                            "anchorXOffset": "0",
                            "anchorYOffset": "0",
                            "anchorIgnoreIfNotPresent": "true",
                            "anchorUnits": "pixels",
                        }],
                    },
                }],
            },
            "status": "sent",
        }
        r = requests.post(
            f"{base_url}/restapi/v2.1/accounts/{account_id}/envelopes",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=envelope,
            timeout=15,
        )
        data = r.json()
        if "envelopeId" not in data:
            return {"error": data.get("message", data.get("errorCode", "Envelope creation failed"))}
        return {
            "envelope_id": data["envelopeId"],
            "status": data.get("status", "sent"),
        }
    except Exception as e:
        return {"error": str(e)}


def docusign_envelope_status(
    access_token: str,
    account_id: str,
    base_url: str,
    envelope_id: str,
) -> dict:
    """Get DocuSign envelope status. Returns {status, signed_at, error?}."""
    try:
        base_url = base_url.rstrip("/")
        r = requests.get(
            f"{base_url}/restapi/v2.1/accounts/{account_id}/envelopes/{envelope_id}",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=15,
        )
        data = r.json()
        if "status" not in data:
            return {"error": data.get("message", "Status fetch failed")}
        return {
            "status": data["status"],
            "signed_at": data.get("completedDateTime", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def docusign_list_envelopes(
    access_token: str,
    account_id: str,
    base_url: str,
) -> list[dict]:
    """List DocuSign envelopes. Returns list of dicts."""
    try:
        base_url = base_url.rstrip("/")
        from_date = (datetime.now(timezone.utc) - timedelta(days=90)).strftime("%Y-%m-%dT%H:%M:%SZ")
        r = requests.get(
            f"{base_url}/restapi/v2.1/accounts/{account_id}/envelopes",
            headers={"Authorization": f"Bearer {access_token}"},
            params={"from_date": from_date, "include": "recipients"},
            timeout=15,
        )
        data = r.json()
        envelopes = data.get("envelopes", [])
        return [
            {
                "envelope_id": e.get("envelopeId", ""),
                "subject": e.get("emailSubject", ""),
                "status": e.get("status", ""),
                "created": e.get("createdDateTime", ""),
                "completed": e.get("completedDateTime", ""),
            }
            for e in envelopes
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# ADOBE SIGN
# ---------------------------------------------------------------------------

_ADOBE_BASE = "https://api.na4.adobesign.com/api/rest/v6"


def adobe_sign_send(
    access_token: str,
    doc_name: str,
    doc_base64: str,
    signer_email: str,
    signer_name: str,
    message: str = "",
) -> dict:
    """Send a document for signature via Adobe Sign. Returns {agreement_id, status, error?}."""
    try:
        # First, get the base URI
        info_r = requests.get(
            "https://api.na4.adobesign.com/api/rest/v6/baseUris",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=15,
        )
        if info_r.status_code == 200:
            base_uri = info_r.json().get("apiAccessPoint", "https://api.na4.adobesign.com/") + "api/rest/v6"
        else:
            base_uri = _ADOBE_BASE

        # Upload transient document
        doc_bytes = base64.b64decode(doc_base64)
        files = {"File": (doc_name, io.BytesIO(doc_bytes), "application/octet-stream")}
        upload_r = requests.post(
            f"{base_uri}/transientDocuments",
            headers={"Authorization": f"Bearer {access_token}"},
            files=files,
            timeout=30,
        )
        upload_data = upload_r.json()
        if "transientDocumentId" not in upload_data:
            return {"error": upload_data.get("message", "Document upload failed")}

        transient_id = upload_data["transientDocumentId"]

        # Create agreement
        agreement = {
            "fileInfos": [{"transientDocumentId": transient_id}],
            "name": doc_name,
            "participantSetsInfo": [{
                "memberInfos": [{"email": signer_email, "name": signer_name}],
                "order": 1,
                "role": "SIGNER",
            }],
            "signatureType": "ESIGN",
            "state": "IN_PROCESS",
            "message": message,
        }
        agree_r = requests.post(
            f"{base_uri}/agreements",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=agreement,
            timeout=15,
        )
        agree_data = agree_r.json()
        if "id" not in agree_data:
            return {"error": agree_data.get("message", "Agreement creation failed")}
        return {
            "agreement_id": agree_data["id"],
            "status": "IN_PROCESS",
        }
    except Exception as e:
        return {"error": str(e)}


def adobe_sign_status(access_token: str, agreement_id: str) -> dict:
    """Get Adobe Sign agreement status."""
    try:
        r = requests.get(
            f"{_ADOBE_BASE}/agreements/{agreement_id}",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=15,
        )
        data = r.json()
        if "status" not in data:
            return {"error": data.get("message", "Status fetch failed")}
        return {"status": data["status"]}
    except Exception as e:
        return {"error": str(e)}


def adobe_sign_list(access_token: str) -> list[dict]:
    """List all Adobe Sign agreements."""
    try:
        r = requests.get(
            f"{_ADOBE_BASE}/agreements",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=15,
        )
        data = r.json()
        items = data.get("userAgreementList", [])
        return [
            {
                "agreement_id": a.get("id", ""),
                "name": a.get("name", ""),
                "status": a.get("status", ""),
                "created": a.get("createdDate", ""),
            }
            for a in items
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# MICROSOFT WORD / DOCX (python-docx)
# ---------------------------------------------------------------------------

def text_to_docx(text: str, title: str = "", author: str = "eLawFirm") -> bytes:
    """Convert plain text to a DOCX file. Returns bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Set author in core properties
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()

        if title:
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        for para_text in text.split("\n"):
            doc.add_paragraph(para_text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        # Return minimal docx bytes with error text
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph(f"Error generating document: {e}")
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            return b""


def markdown_to_docx(md_text: str, title: str = "", author: str = "eLawFirm") -> bytes:
    """Convert markdown text to a DOCX file. Returns bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()

        if title:
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        lines = md_text.split("\n")
        for line in lines:
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("**") and line.endswith("**") and len(line) > 4:
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip() == "---":
                doc.add_paragraph("_" * 60)
            else:
                p = doc.add_paragraph()
                # Handle inline bold
                parts = re.split(r"\*\*(.*?)\*\*", line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        return text_to_docx(md_text, title=title, author=author)


# ---------------------------------------------------------------------------
# CALENDAR — ICS + Google Calendar API + Outlook Graph
# ---------------------------------------------------------------------------

def make_ics(events: list[dict]) -> bytes:
    """Generate an ICS calendar file from a list of event dicts.
    Each event: {title, start, end, description, location}
    start/end: ISO 8601 strings e.g. "2025-06-01T10:00:00"
    """
    try:
        lines = [
            "BEGIN:VCALENDAR",
            "VERSION:2.0",
            "PRODID:-//eLawFirm//EN",
            "CALSCALE:GREGORIAN",
            "METHOD:PUBLISH",
        ]
        for ev in events:
            uid = str(uuid.uuid4())
            start = ev.get("start", "")
            end = ev.get("end", "")

            def _fmt_dt(dt_str: str) -> str:
                try:
                    dt = datetime.fromisoformat(dt_str.replace("Z", "+00:00"))
                    return dt.strftime("%Y%m%dT%H%M%SZ")
                except Exception:
                    return dt_str

            dtstart = _fmt_dt(start)
            dtend = _fmt_dt(end) if end else dtstart
            title = (ev.get("title", "Event") or "").replace("\n", " ")
            desc = (ev.get("description", "") or "").replace("\n", "\\n")
            loc = (ev.get("location", "") or "").replace("\n", " ")

            lines += [
                "BEGIN:VEVENT",
                f"UID:{uid}",
                f"DTSTART:{dtstart}",
                f"DTEND:{dtend}",
                f"SUMMARY:{title}",
                f"DESCRIPTION:{desc}",
                f"LOCATION:{loc}",
                "END:VEVENT",
            ]
        lines.append("END:VCALENDAR")
        return "\r\n".join(lines).encode("utf-8")
    except Exception as e:
        return f"BEGIN:VCALENDAR\r\nVERSION:2.0\r\n# Error: {e}\r\nEND:VCALENDAR".encode()


def google_calendar_create(
    access_token: str,
    title: str,
    start_iso: str,
    end_iso: str,
    description: str = "",
    location: str = "",
    calendar_id: str = "primary",
) -> dict:
    """Create a Google Calendar event. Returns {id, url, error?}."""
    try:
        event = {
            "summary": title,
            "description": description,
            "location": location,
            "start": {"dateTime": start_iso, "timeZone": "UTC"},
            "end": {"dateTime": end_iso, "timeZone": "UTC"},
        }
        r = requests.post(
            f"https://www.googleapis.com/calendar/v3/calendars/{quote(calendar_id)}/events",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=event,
            timeout=15,
        )
        data = r.json()
        if "id" not in data:
            return {"error": data.get("error", {}).get("message", "Event creation failed")}
        return {
            "id": data["id"],
            "url": data.get("htmlLink", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def google_calendar_list_events(
    access_token: str,
    days_ahead: int = 14,
) -> list[dict]:
    """List upcoming Google Calendar events."""
    try:
        now = datetime.now(timezone.utc).isoformat()
        future = (datetime.now(timezone.utc) + timedelta(days=days_ahead)).isoformat()
        r = requests.get(
            "https://www.googleapis.com/calendar/v3/calendars/primary/events",
            headers={"Authorization": f"Bearer {access_token}"},
            params={
                "timeMin": now,
                "timeMax": future,
                "singleEvents": "true",
                "orderBy": "startTime",
                "maxResults": 50,
            },
            timeout=15,
        )
        data = r.json()
        if "items" not in data:
            return [{"error": data.get("error", {}).get("message", "List failed")}]
        return [
            {
                "id": e.get("id", ""),
                "title": e.get("summary", ""),
                "start": e.get("start", {}).get("dateTime", e.get("start", {}).get("date", "")),
                "end": e.get("end", {}).get("dateTime", e.get("end", {}).get("date", "")),
                "location": e.get("location", ""),
                "description": e.get("description", ""),
            }
            for e in data["items"]
        ]
    except Exception as e:
        return [{"error": str(e)}]


def outlook_calendar_create(
    access_token: str,
    title: str,
    start_iso: str,
    end_iso: str,
    description: str = "",
    location: str = "",
) -> dict:
    """Create an Outlook Calendar event via Microsoft Graph. Returns {id, url, error?}."""
    try:
        event = {
            "subject": title,
            "body": {"contentType": "Text", "content": description},
            "start": {"dateTime": start_iso, "timeZone": "UTC"},
            "end": {"dateTime": end_iso, "timeZone": "UTC"},
            "location": {"displayName": location},
        }
        r = requests.post(
            f"{_GRAPH_API}/me/events",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=event,
            timeout=15,
        )
        data = r.json()
        if "id" not in data:
            return {"error": data.get("error", {}).get("message", "Event creation failed")}
        return {
            "id": data["id"],
            "url": data.get("webLink", ""),
        }
    except Exception as e:
        return {"error": str(e)}


# ---------------------------------------------------------------------------
# WESTLAW
# ---------------------------------------------------------------------------

def westlaw_search(
    api_key: str,
    query: str,
    jurisdiction: str = "",
    max_results: int = 10,
) -> list[dict]:
    """Search Westlaw for legal cases/statutes. Returns list of {title, citation, excerpt, url}."""
    if not api_key:
        return [{"error": "API key required"}]
    try:
        payload: dict[str, Any] = {
            "query": query,
            "limit": max_results,
        }
        if jurisdiction:
            payload["jurisdiction"] = jurisdiction

        r = requests.post(
            "https://api.westlaw.com/v1/search",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=payload,
            timeout=15,
        )
        if r.status_code == 401:
            return [{"error": "Invalid API key"}]
        if r.status_code == 403:
            return [{"error": "API key does not have access to this resource"}]
        data = r.json()
        results = data.get("results", data.get("documents", []))
        return [
            {
                "title": item.get("title", item.get("name", "")),
                "citation": item.get("citation", item.get("cite", "")),
                "excerpt": item.get("excerpt", item.get("snippet", item.get("summary", ""))),
                "url": item.get("url", item.get("documentUrl", "")),
            }
            for item in results
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# LEXISNEXIS
# ---------------------------------------------------------------------------

def lexisnexis_search(
    api_key: str,
    query: str,
    jurisdiction: str = "",
    max_results: int = 10,
) -> list[dict]:
    """Search LexisNexis for legal cases/statutes. Returns list of {title, citation, excerpt, url}."""
    if not api_key:
        return [{"error": "API key required"}]
    try:
        params: dict[str, Any] = {
            "q": query,
            "pageSize": max_results,
        }
        if jurisdiction:
            params["jurisdiction"] = jurisdiction

        r = requests.get(
            "https://services.lexisnexis.com/v1/documents",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Accept": "application/json",
            },
            params=params,
            timeout=15,
        )
        if r.status_code == 401:
            return [{"error": "Invalid API key"}]
        if r.status_code == 403:
            return [{"error": "API key does not have access to this resource"}]
        data = r.json()
        results = data.get("documents", data.get("results", []))
        return [
            {
                "title": item.get("title", item.get("name", "")),
                "citation": item.get("citation", item.get("normalizedCite", "")),
                "excerpt": item.get("excerpt", item.get("snippet", item.get("summary", ""))),
                "url": item.get("url", item.get("documentUrl", "")),
            }
            for item in results
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# XERO (OAuth2)
# ---------------------------------------------------------------------------

_XERO_AUTH_URL = "https://login.xero.com/identity/connect/authorize"
_XERO_TOKEN_URL = "https://identity.xero.com/connect/token"
_XERO_API = "https://api.xero.com/api.xro/2.0"


def xero_auth_url(
    client_id: str,
    redirect_uri: str = "https://app.elaw.firm/callback",
) -> str:
    """Return the Xero OAuth2 authorization URL."""
    try:
        params = {
            "response_type": "code",
            "client_id": client_id,
            "redirect_uri": redirect_uri,
            "scope": "accounting.transactions accounting.contacts offline_access",
            "state": str(uuid.uuid4()),
        }
        return f"{_XERO_AUTH_URL}?{urlencode(params)}"
    except Exception as e:
        return f"error:{e}"


def xero_exchange_code(
    client_id: str,
    client_secret: str,
    code: str,
    redirect_uri: str,
) -> dict:
    """Exchange Xero auth code for tokens. Returns {access_token, refresh_token, error?}."""
    try:
        creds = base64.b64encode(f"{client_id}:{client_secret}".encode()).decode()
        r = requests.post(
            _XERO_TOKEN_URL,
            headers={
                "Authorization": f"Basic {creds}",
                "Content-Type": "application/x-www-form-urlencoded",
            },
            data={
                "grant_type": "authorization_code",
                "code": code,
                "redirect_uri": redirect_uri,
            },
            timeout=15,
        )
        data = r.json()
        if "access_token" not in data:
            return {"error": data.get("error_description", data.get("error", "Token exchange failed"))}
        return {
            "access_token": data["access_token"],
            "refresh_token": data.get("refresh_token", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def xero_get_tenants(access_token: str) -> list[dict]:
    """Get Xero tenants (organisations). Returns list of {tenantId, tenantName}."""
    try:
        r = requests.get(
            "https://api.xero.com/connections",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=15,
        )
        data = r.json()
        if isinstance(data, list):
            return [
                {
                    "tenantId": t.get("tenantId", ""),
                    "tenantName": t.get("tenantName", ""),
                }
                for t in data
            ]
        return [{"error": str(data)}]
    except Exception as e:
        return [{"error": str(e)}]


def xero_create_invoice(
    access_token: str,
    tenant_id: str,
    contact_name: str,
    line_items: list[dict],
    due_date: str,
    invoice_number: str,
) -> dict:
    """Create a Xero invoice. Returns {invoice_id, invoice_number, status, url, error?}."""
    try:
        xero_line_items = []
        for li in line_items:
            xero_line_items.append({
                "Description": li.get("description", ""),
                "Quantity": li.get("quantity", 1),
                "UnitAmount": li.get("unit_amount", 0),
                "AccountCode": li.get("account_code", "200"),
            })

        invoice = {
            "Type": "ACCREC",
            "Contact": {"Name": contact_name},
            "LineItems": xero_line_items,
            "DueDate": due_date,
            "InvoiceNumber": invoice_number,
            "Status": "DRAFT",
        }
        r = requests.post(
            f"{_XERO_API}/Invoices",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Xero-tenant-id": tenant_id,
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json={"Invoices": [invoice]},
            timeout=15,
        )
        data = r.json()
        invoices = data.get("Invoices", [])
        if not invoices:
            return {"error": data.get("Detail", data.get("Message", "Invoice creation failed"))}
        inv = invoices[0]
        return {
            "invoice_id": inv.get("InvoiceID", ""),
            "invoice_number": inv.get("InvoiceNumber", invoice_number),
            "status": inv.get("Status", "DRAFT"),
            "url": f"https://go.xero.com/AccountsReceivable/View.aspx?InvoiceID={inv.get('InvoiceID', '')}",
        }
    except Exception as e:
        return {"error": str(e)}


def xero_list_invoices(access_token: str, tenant_id: str) -> list[dict]:
    """List Xero invoices."""
    try:
        r = requests.get(
            f"{_XERO_API}/Invoices",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Xero-tenant-id": tenant_id,
                "Accept": "application/json",
            },
            params={"where": "Type==\"ACCREC\"", "order": "Date DESC"},
            timeout=15,
        )
        data = r.json()
        return [
            {
                "invoice_id": i.get("InvoiceID", ""),
                "invoice_number": i.get("InvoiceNumber", ""),
                "contact": i.get("Contact", {}).get("Name", ""),
                "amount": i.get("Total", 0),
                "status": i.get("Status", ""),
                "due_date": i.get("DueDate", ""),
            }
            for i in data.get("Invoices", [])
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# QUICKBOOKS ONLINE (OAuth2)
# ---------------------------------------------------------------------------

_QBO_AUTH_URL = "https://appcenter.intuit.com/connect/oauth2"
_QBO_TOKEN_URL = "https://oauth.platform.intuit.com/oauth2/v1/tokens/bearer"
_QBO_SCOPE = "com.intuit.quickbooks.accounting"


def qbo_auth_url(client_id: str, redirect_uri: str) -> str:
    """Return the QuickBooks OAuth2 authorization URL."""
    try:
        params = {
            "client_id": client_id,
            "scope": _QBO_SCOPE,
            "redirect_uri": redirect_uri,
            "response_type": "code",
            "state": str(uuid.uuid4()),
        }
        return f"{_QBO_AUTH_URL}?{urlencode(params)}"
    except Exception as e:
        return f"error:{e}"


def qbo_exchange_code(
    client_id: str,
    client_secret: str,
    code: str,
    redirect_uri: str,
) -> dict:
    """Exchange QBO auth code for tokens. Returns {access_token, refresh_token, error?}."""
    try:
        creds = base64.b64encode(f"{client_id}:{client_secret}".encode()).decode()
        r = requests.post(
            _QBO_TOKEN_URL,
            headers={
                "Authorization": f"Basic {creds}",
                "Content-Type": "application/x-www-form-urlencoded",
                "Accept": "application/json",
            },
            data={
                "grant_type": "authorization_code",
                "code": code,
                "redirect_uri": redirect_uri,
            },
            timeout=15,
        )
        data = r.json()
        if "access_token" not in data:
            return {"error": data.get("error_description", data.get("error", "Token exchange failed"))}
        return {
            "access_token": data["access_token"],
            "refresh_token": data.get("refresh_token", ""),
        }
    except Exception as e:
        return {"error": str(e)}


def qbo_create_invoice(
    access_token: str,
    realm_id: str,
    customer_name: str,
    line_items: list[dict],
    due_date: str,
    invoice_number: str,
) -> dict:
    """Create a QuickBooks invoice. Returns {invoice_id, error?}."""
    try:
        base_url = f"https://quickbooks.api.intuit.com/v3/company/{realm_id}"

        # Find or create customer
        customer_query = requests.get(
            f"{base_url}/query",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Accept": "application/json",
            },
            params={"query": f"select * from Customer WHERE DisplayName = '{customer_name}' MAXRESULTS 1"},
            timeout=15,
        )
        customers = customer_query.json().get("QueryResponse", {}).get("Customer", [])
        if customers:
            customer_ref = {"value": customers[0]["Id"], "name": customer_name}
        else:
            create_r = requests.post(
                f"{base_url}/customer",
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                },
                json={"DisplayName": customer_name},
                timeout=15,
            )
            created = create_r.json().get("Customer", {})
            customer_ref = {"value": created.get("Id", "1"), "name": customer_name}

        qbo_lines = []
        for i, li in enumerate(line_items):
            qbo_lines.append({
                "Amount": float(li.get("quantity", 1)) * float(li.get("unit_amount", 0)),
                "DetailType": "SalesItemLineDetail",
                "Description": li.get("description", ""),
                "SalesItemLineDetail": {
                    "Qty": li.get("quantity", 1),
                    "UnitPrice": li.get("unit_amount", 0),
                },
            })

        invoice = {
            "Line": qbo_lines,
            "CustomerRef": customer_ref,
            "DueDate": due_date,
            "DocNumber": invoice_number,
        }
        r = requests.post(
            f"{base_url}/invoice",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
                "Accept": "application/json",
            },
            json=invoice,
            timeout=15,
        )
        data = r.json()
        inv = data.get("Invoice", {})
        if not inv:
            return {"error": data.get("Fault", {}).get("Error", [{}])[0].get("Detail", "Invoice creation failed")}
        return {"invoice_id": inv.get("Id", "")}
    except Exception as e:
        return {"error": str(e)}


def qbo_list_invoices(access_token: str, realm_id: str) -> list[dict]:
    """List QuickBooks invoices."""
    try:
        base_url = f"https://quickbooks.api.intuit.com/v3/company/{realm_id}"
        r = requests.get(
            f"{base_url}/query",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Accept": "application/json",
            },
            params={"query": "select * from Invoice ORDERBY MetaData.CreateTime DESC MAXRESULTS 50"},
            timeout=15,
        )
        data = r.json()
        invoices = data.get("QueryResponse", {}).get("Invoice", [])
        return [
            {
                "invoice_id": inv.get("Id", ""),
                "doc_number": inv.get("DocNumber", ""),
                "customer": inv.get("CustomerRef", {}).get("name", ""),
                "amount": inv.get("TotalAmt", 0),
                "balance": inv.get("Balance", 0),
                "due_date": inv.get("DueDate", ""),
            }
            for inv in invoices
        ]
    except Exception as e:
        return [{"error": str(e)}]


# ---------------------------------------------------------------------------
# ZAPIER / MAKE WEBHOOKS
# ---------------------------------------------------------------------------

def zapier_trigger(webhook_url: str, event_type: str, payload: dict) -> dict:
    """Trigger a Zapier webhook. Returns {success, status_code, error?}."""
    try:
        body = {"event": event_type, "data": payload, "source": "eLawFirm"}
        r = requests.post(webhook_url, json=body, timeout=15)
        return {
            "success": r.status_code in (200, 201),
            "status_code": r.status_code,
        }
    except Exception as e:
        return {"success": False, "status_code": 0, "error": str(e)}


def make_trigger(webhook_url: str, payload: dict) -> dict:
    """Trigger a Make (Integromat) webhook. Returns {success, status_code, error?}."""
    try:
        r = requests.post(webhook_url, json=payload, timeout=15)
        return {
            "success": r.status_code in (200, 201, 204),
            "status_code": r.status_code,
        }
    except Exception as e:
        return {"success": False, "status_code": 0, "error": str(e)}


def webhook_test(url: str, payload: dict | None = None) -> dict:
    """Test a webhook URL. Returns {success, status_code, latency_ms, error?}."""
    try:
        test_payload = payload or {"test": True, "source": "eLawFirm", "timestamp": datetime.now(timezone.utc).isoformat()}
        start = time.monotonic()
        r = requests.post(url, json=test_payload, timeout=15)
        latency_ms = int((time.monotonic() - start) * 1000)
        return {
            "success": r.status_code in (200, 201, 204),
            "status_code": r.status_code,
            "latency_ms": latency_ms,
        }
    except Exception as e:
        return {"success": False, "status_code": 0, "latency_ms": 0, "error": str(e)}


# ---------------------------------------------------------------------------
# SSO / SAML 2.0
# ---------------------------------------------------------------------------

def saml_generate_auth_request(
    idp_sso_url: str,
    sp_entity_id: str,
    acs_url: str,
) -> dict:
    """Generate a SAML 2.0 AuthnRequest and return redirect URL + request ID."""
    try:
        request_id = "_" + uuid.uuid4().hex
        issue_instant = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        auth_request = (
            f'<samlp:AuthnRequest'
            f' xmlns:samlp="urn:oasis:names:tc:SAML:2.0:protocol"'
            f' xmlns:saml="urn:oasis:names:tc:SAML:2.0:assertion"'
            f' ID="{request_id}"'
            f' Version="2.0"'
            f' IssueInstant="{issue_instant}"'
            f' AssertionConsumerServiceURL="{acs_url}"'
            f' ProtocolBinding="urn:oasis:names:tc:SAML:2.0:bindings:HTTP-POST">'
            f'<saml:Issuer>{sp_entity_id}</saml:Issuer>'
            f'<samlp:NameIDPolicy'
            f' Format="urn:oasis:names:tc:SAML:1.1:nameid-format:emailAddress"'
            f' AllowCreate="true"/>'
            f'</samlp:AuthnRequest>'
        )

        encoded = base64.b64encode(auth_request.encode()).decode()
        params = {"SAMLRequest": encoded, "RelayState": request_id}
        redirect_url = f"{idp_sso_url}?{urlencode(params)}"
        return {"redirect_url": redirect_url, "request_id": request_id}
    except Exception as e:
        return {"error": str(e), "redirect_url": "", "request_id": ""}


def saml_parse_response(saml_response_b64: str, idp_cert: str) -> dict:
    """Parse a SAML 2.0 response (base64-encoded). Returns {email, name, attributes, error?}."""
    try:
        xml_bytes = base64.b64decode(saml_response_b64)
        root = ET.fromstring(xml_bytes)

        ns = {
            "saml": "urn:oasis:names:tc:SAML:2.0:assertion",
            "samlp": "urn:oasis:names:tc:SAML:2.0:protocol",
        }

        # Check status
        status_code = root.find(".//samlp:StatusCode", ns)
        if status_code is not None:
            val = status_code.attrib.get("Value", "")
            if "Success" not in val:
                return {"error": f"SAML authentication failed: {val}", "email": "", "name": "", "attributes": {}}

        # Extract NameID (email)
        name_id_el = root.find(".//saml:NameID", ns)
        email = name_id_el.text.strip() if name_id_el is not None and name_id_el.text else ""

        # Extract attributes
        attributes: dict[str, str] = {}
        for attr in root.findall(".//saml:Attribute", ns):
            attr_name = attr.attrib.get("Name", "")
            values = [v.text for v in attr.findall("saml:AttributeValue", ns) if v.text]
            if attr_name and values:
                attributes[attr_name] = values[0]

        name = (
            attributes.get("http://schemas.xmlsoap.org/ws/2005/05/identity/claims/name")
            or attributes.get("displayName")
            or attributes.get("cn")
            or attributes.get("name")
            or email
        )

        return {"email": email, "name": name, "attributes": attributes}
    except Exception as e:
        return {"error": str(e), "email": "", "name": "", "attributes": {}}


def saml_generate_metadata(
    entity_id: str,
    acs_url: str,
    sp_cert: str = "",
) -> str:
    """Generate SAML SP metadata XML."""
    try:
        cert_block = ""
        if sp_cert:
            clean_cert = sp_cert.replace("-----BEGIN CERTIFICATE-----", "").replace("-----END CERTIFICATE-----", "").strip()
            cert_block = (
                f'<md:KeyDescriptor use="signing">'
                f'<ds:KeyInfo xmlns:ds="http://www.w3.org/2000/09/xmldsig#">'
                f'<ds:X509Data><ds:X509Certificate>{clean_cert}</ds:X509Certificate></ds:X509Data>'
                f'</ds:KeyInfo></md:KeyDescriptor>'
                f'<md:KeyDescriptor use="encryption">'
                f'<ds:KeyInfo xmlns:ds="http://www.w3.org/2000/09/xmldsig#">'
                f'<ds:X509Data><ds:X509Certificate>{clean_cert}</ds:X509Certificate></ds:X509Data>'
                f'</ds:KeyInfo></md:KeyDescriptor>'
            )

        metadata = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<md:EntityDescriptor'
            ' xmlns:md="urn:oasis:names:tc:SAML:2.0:metadata"'
            f' entityID="{entity_id}">'
            '<md:SPSSODescriptor'
            ' AuthnRequestsSigned="false"'
            ' WantAssertionsSigned="true"'
            ' protocolSupportEnumeration="urn:oasis:names:tc:SAML:2.0:protocol">'
            f'{cert_block}'
            '<md:NameIDFormat>'
            'urn:oasis:names:tc:SAML:1.1:nameid-format:emailAddress'
            '</md:NameIDFormat>'
            f'<md:AssertionConsumerService'
            f' Binding="urn:oasis:names:tc:SAML:2.0:bindings:HTTP-POST"'
            f' Location="{acs_url}"'
            f' index="1"/>'
            '</md:SPSSODescriptor>'
            '</md:EntityDescriptor>'
        )
        return metadata
    except Exception as e:
        return f"<!-- Error generating metadata: {e} -->"
